# Core/images_figures.py
# -*- coding: utf-8 -*-
import re
import io
import tempfile
from typing import List, Tuple, Dict, Optional
from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx import Document
from docx.text.paragraph import Paragraph
from docx.shared import Emu
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image  # pour valider / redimensionner les images

try:
    import fitz  # PyMuPDF pour PDF
except Exception:
    fitz = None  # type: ignore

# ==========================
# Utils espace / pattern
# ==========================

def _iter_all_paragraphs_doc(doc: Document):
    """
    Retourne tous les Paragraph du document, y compris ceux des tableaux,
    en-têtes et pieds de page.
    """
    def iter_container(container):
        for p in getattr(container, "paragraphs", []):
            yield p
        for table in getattr(container, "tables", []):
            for row in table.rows:
                for cell in row.cells:
                    yield from iter_container(cell)

    # Corps principal
    yield from iter_container(doc)

    # Headers / footers
    for section in doc.sections:
        yield from iter_container(section.header)
        yield from iter_container(section.footer)

FIG_PATTERN = re.compile(r"(?i)\b(?:cf\.\s*)?(?:fig(?:\.|ure)?|image)\s*[:\-–]?\s*(\d+)\b")

def normalize_and_filter_images(
    images: List[Tuple[bytes, str]],
    *,
    min_side: int = 400,
    max_images: int = 30,
) -> List[Tuple[bytes, str]]:
    """
    - Ouvre chaque image avec Pillow pour vérifier qu'elle est valide.
    - Ne garde que les images dont la plus grande dimension >= min_side.
    - Réencode en PNG (format sûr pour python-docx).
    - Trie par taille (aire w*h décroissante) et garde max_images.

    Retourne une liste [(bytes_png, "png"), ...].
    """
    cleaned: List[Tuple[bytes, str, int]] = []

    for idx, (data, ext) in enumerate(images, start=1):
        if not data:
            continue
        try:
            im = Image.open(io.BytesIO(data))
            im.load()  # force le décodage
            w, h = im.size
        except Exception:
            # image illisible -> on passe
            continue

        if max(w, h) < min_side:
            # trop petite (icône, logo minuscule, etc.) -> on ignore
            continue

        # réencodage en PNG propre
        try:
            out = io.BytesIO()
            im.save(out, format="PNG")
            png_bytes = out.getvalue()
        except Exception:
            continue

        area = w * h
        cleaned.append((png_bytes, "png", area))

    # tri par taille décroissante
    cleaned.sort(key=lambda x: x[2], reverse=True)
    # on garde au maximum max_images
    selected = cleaned[:max_images]

    print(
        f"[images_figures] normalize_and_filter_images: "
        f"{len(images)} image(s) en entrée, {len(selected)} retenue(s) "
        f"(min_side={min_side}, max_images={max_images})"
    )

    # on renvoie sans l'aire
    return [(b, ext) for (b, ext, _area) in selected]

def _norm_space(s: str) -> str:
    # normalise NBSP & NNBSP en espace simple
    return (s or "").replace("\u00A0", " ").replace("\u202F", " ")


# ==========================
# Extraction images DOCX source (ordre apparition) : a:blip + VML imagedata
# ==========================

def _attr_local(elm, local_name: str) -> Optional[str]:
    """Retourne la valeur d'un attribut dont le *local-name()* = local_name, sinon None."""
    for k, v in elm.attrib.items():
        if k.split("}")[-1] == local_name:
            return v
    return None


def _collect_embeds_live(parent_elm) -> List[str]:
    """
    Sur un élément OXML 'vivant' (doc._element, header._element...), collecte
    les rId d'images dans l'ordre d'apparition :
      - DrawingML:  //a:blip  -> @r:embed | @r:link
      - VML:       //w:pict//v:imagedata -> @r:id | @r:link
    """
    rids: List[str] = []
    # 1) a:blip
    for blip in parent_elm.xpath(".//*[local-name()='blip']"):
        rid = _attr_local(blip, "embed") or _attr_local(blip, "link")
        if rid:
            rids.append(rid)
    # 2) v:imagedata sous w:pict
    for imd in parent_elm.xpath(".//*[local-name()='pict']//*[local-name()='imagedata']"):
        rid = _attr_local(imd, "id") or _attr_local(imd, "link")
        if rid:
            rids.append(rid)
    return rids


def _resolve_rids(part, rids: List[str]) -> List[Tuple[bytes, str]]:
    """
    Convertit une liste de rId en [(bytes, ext)] en résolvant part.rels.
    """
    out: List[Tuple[bytes, str]] = []
    for rid in rids:
        rel = part.rels.get(rid)
        if not rel:
            continue
        target = rel._target
        data = getattr(target, "blob", None)
        if not data:
            continue
        partname = str(getattr(target, "partname", ""))  # e.g. /word/media/image1.png
        ext = partname.rsplit(".", 1)[-1].lower() if "." in partname else "png"
        out.append((data, ext))
    return out


def extract_images_from_docx_ordered_any(path: str, include_header_footer: bool = False) -> List[Tuple[bytes, str]]:
    """
    Extrait toutes les images d'un DOCX dans l'ordre d'apparition (a:blip + VML imagedata).
    Retourne [(data, ext)].
    """
    import docx
    doc = docx.Document(path)
    images: List[Tuple[bytes, str]] = []
    images.extend(_resolve_rids(doc.part, _collect_embeds_live(doc._element)))
    if include_header_footer:
        for section in doc.sections:
            if section.header:
                images.extend(_resolve_rids(section.header.part, _collect_embeds_live(section.header._element)))
            if section.footer:
                images.extend(_resolve_rids(section.footer.part, _collect_embeds_live(section.footer._element)))
    print(f"[images_figures] extract_images_from_docx_ordered_any: {len(images)} image(s) trouvée(s) dans {path}")
    return images


def extract_images_from_docx_bytes(data: bytes) -> List[Tuple[bytes, str]]:
    """
    Helper : écrit les bytes DOCX dans un fichier temp et appelle extract_images_from_docx_ordered_any.
    """
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    tmp.write(data)
    tmp.close()
    return extract_images_from_docx_ordered_any(tmp.name, include_header_footer=False)


def extract_images_from_pdf_bytes(data: bytes) -> List[Tuple[bytes, str]]:
    """
    Extrait les images d'un PDF (bytes) via PyMuPDF.
    Retourne [(bytes, ext)].
    """
    if not fitz:
        print("[images_figures] PyMuPDF (fitz) non disponible, impossible d'extraire les images PDF.")
        return []
    images: List[Tuple[bytes, str]] = []
    try:
        with fitz.open(stream=data, filetype="pdf") as doc:
            for page in doc:
                for xref, *_ in page.get_images(full=True):
                    base = doc.extract_image(xref)
                    img = base.get("image")
                    ext = (base.get("ext") or "png").lower()
                    if img:
                        images.append((img, ext))
    except Exception as e:
        print(f"[images_figures] Erreur extract_images_from_pdf_bytes: {e}")
    print(f"[images_figures] extract_images_from_pdf_bytes: {len(images)} image(s) trouvée(s) dans le PDF")
    return images


# ==========================
# Construction d'un DOCX "source d'images"
# ==========================

def build_docx_from_images(images: List[Tuple[bytes, str]]) -> str:
    """
    Construit un DOCX temporaire qui contient toutes les images fournies (une par paragraphe).
    Les images sont supposées déjà filtrées/normalisées (ex: via normalize_and_filter_images).
    """
    doc = Document()
    count_ok = 0
    for i, (img_bytes, _ext) in enumerate(images, start=1):
        try:
            p = doc.add_paragraph()
            run = p.add_run()
            run.add_picture(io.BytesIO(img_bytes))
            count_ok += 1
        except Exception as e:
            # on log et on ignore l'image problématique
            print(f"[images_figures] build_docx_from_images: image {i} ignorée (erreur add_picture: {e})")

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)
    print(
        f"[images_figures] build_docx_from_images: DOCX temporaire {tmp.name} "
        f"avec {count_ok} image(s) insérée(s) sur {len(images)}"
    )
    return tmp.name



# ==========================
# Arbre vivant : trouver le w:p hôte dans le body, insérer un w:p, y ajouter une image
# ==========================

def _find_body_host_paragraph_live(w_t) -> Optional[object]:
    """
    Depuis un w:t (élément vivant), remonte au w:p qui appartient au w:body.
    Si le w:t est dans une textbox, on remonte au w:p externe hébergé par le body.
    """
    # remonter au w:p le plus proche
    el = w_t
    while el is not None and el.tag != qn("w:p"):
        el = el.getparent()
    if el is None:
        return None

    # si ce w:p appartient au body, ok
    anc = el
    while anc is not None and anc.tag != qn("w:body"):
        anc = anc.getparent()
    if anc is not None and anc.tag == qn("w:body"):
        return el

    # sinon, remonter plus haut jusqu'à trouver un w:p ayant un ancêtre w:body
    p_candidate = el.getparent()
    while p_candidate is not None:
        # chercher le prochain ancêtre w:p
        while p_candidate is not None and p_candidate.tag != qn("w:p"):
            p_candidate = p_candidate.getparent()
        if p_candidate is None:
            break
        anc = p_candidate
        while anc is not None and anc.tag != qn("w:body"):
            anc = anc.getparent()
        if anc is not None and anc.tag == qn("w:body"):
            return p_candidate
        p_candidate = p_candidate.getparent()
    return None


def _insert_paragraph_after_live(p_body):
    """
    Insère un nouveau <w:p> juste après p_body (w:p du corps vivant) et le retourne.
    """
    new_p = OxmlElement("w:p")
    p_body.addnext(new_p)
    return new_p


def _add_picture_to_paragraph_live(doc: Document, p_elm, image_bytes: bytes, width_emu: int):
    """
    Ajoute une image dans le w:p fourni (arbre vivant) via l'API python-docx.
    """
    para = Paragraph(p_elm, doc._body)  # doc._body est un BlockItemContainer valide
    run = para.add_run()
    run.add_picture(io.BytesIO(image_bytes), width=Emu(width_emu))


def _max_text_width_emu(doc: Document) -> int:
    """
    Calcule la largeur disponible pour une image INLINE.
    - Si le document est en colonnes (ex: 2), renvoie la largeur d'une colonne.
    - Sinon, renvoie la largeur pleine (page - marges).
    """
    sect = doc.sections[0]
    full_text_width = int(sect.page_width - sect.left_margin - sect.right_margin)

    # Détection colonnes via sectPr (python-docx ne l’expose pas proprement)
    try:
        cols_elms = sect._sectPr.xpath(".//*[local-name()='cols']")
        if cols_elms:
            cols = cols_elms[0]
            num = cols.get(qn("w:num")) or cols.get("num")  # selon version
            space = cols.get(qn("w:space")) or cols.get("space")
            num_cols = int(num) if num and str(num).isdigit() else 1
            space_twips = int(space) if space and str(space).isdigit() else 0

            # Conversion Twips -> EMU (1 twip = 635 EMU)
            space_emu = space_twips * 635

            if num_cols and num_cols > 1:
                # largeur colonne ≈ (largeur totale - espaces entre colonnes) / nb colonnes
                total_gutters = space_emu * (num_cols - 1)
                col_width = max(1, (full_text_width - total_gutters) // num_cols)
                return int(col_width)
    except Exception:
        pass

    return int(full_text_width)



# ==========================
# Champs Word (SEQ, REF, Bookmarks) pour vraies légendes
# ==========================

def _create_seq_field(seq_name: str = "Figure", initial_value: int = 1) -> OxmlElement:
    """
    Crée un champ SEQ Word pour la numérotation automatique.
    Génère: <w:fldSimple w:instr=" SEQ Figure \* ARABIC "><w:r><w:t>N</w:t></w:r></w:fldSimple>

    initial_value: Valeur affichée par défaut (sera recalculée par Word avec F9)
    """
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), f" SEQ {seq_name} \\* ARABIC ")
    # Valeur par défaut (sera mise à jour par Word)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = str(initial_value)
    r.append(t)
    fld.append(r)
    return fld


def _create_bookmark_start(bookmark_id: int, bookmark_name: str) -> OxmlElement:
    """Crée un élément w:bookmarkStart."""
    bm_start = OxmlElement("w:bookmarkStart")
    bm_start.set(qn("w:id"), str(bookmark_id))
    bm_start.set(qn("w:name"), bookmark_name)
    return bm_start


def _create_bookmark_end(bookmark_id: int) -> OxmlElement:
    """Crée un élément w:bookmarkEnd."""
    bm_end = OxmlElement("w:bookmarkEnd")
    bm_end.set(qn("w:id"), str(bookmark_id))
    return bm_end


def _create_ref_field(bookmark_name: str, caption_label: str = "Figure", display_num: int = 1) -> OxmlElement:
    """
    Crée un champ REF qui pointe vers un signet.
    Le flag \h crée un hyperlien.

    display_num: Numéro affiché par défaut (sera recalculé par Word avec F9)
    """
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), f" REF {bookmark_name} \\h ")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = f"{caption_label} {display_num}"
    r.append(t)
    fld.append(r)
    return fld


def _get_next_bookmark_id(doc: Document) -> int:
    """
    Trouve le prochain ID de signet disponible dans le document.
    Les IDs doivent être uniques dans tout le document.
    """
    max_id = 0
    try:
        for bm in doc._element.xpath("//w:bookmarkStart"):
            try:
                bm_id = int(bm.get(qn("w:id")) or 0)
                if bm_id > max_id:
                    max_id = bm_id
            except (ValueError, TypeError):
                pass
    except Exception:
        pass
    return max_id + 1


def _add_caption_with_seq(
    doc: Document,
    after_p_elm,
    fig_num: int,
    caption_text: str,
    caption_label: str = "Figure",
    caption_style: str = "Caption",
    nbsp: str = "\u202F",
    seq_name: str = "Figure",
    seq_order: int = None,
) -> tuple:
    """
    Ajoute une légende avec un vrai champ SEQ et un signet pour les références croisées.

    Structure générée:
    <w:p>
      <w:bookmarkStart w:id="X" w:name="_Ref_Figure_N"/>
      <w:r><w:t>Figure </w:t></w:r>
      <w:fldSimple w:instr=" SEQ Figure \* ARABIC ">...</w:fldSimple>
      <w:bookmarkEnd w:id="X"/>
      <w:r><w:t> : Légende</w:t></w:r>
    </w:p>

    seq_order: Numéro d'ordre réel d'insertion (utilisé pour la valeur par défaut du SEQ)

    Retourne (bookmark_name, bookmark_id, seq_order) pour les références croisées.
    """
    # Utiliser seq_order si fourni, sinon fig_num
    display_num = seq_order if seq_order is not None else fig_num

    # Créer le paragraphe de légende
    cap_p = _insert_paragraph_after_live(after_p_elm)
    para_cap = Paragraph(cap_p, doc._body)

    # Appliquer le style
    try:
        if caption_style:
            para_cap.style = caption_style
    except Exception:
        pass

    para_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Générer un nom de signet unique (basé sur fig_num pour le mapping)
    bookmark_name = f"_Ref_{seq_name}_{fig_num}"
    bookmark_id = _get_next_bookmark_id(doc)

    # 1. Ajouter le bookmarkStart
    bm_start = _create_bookmark_start(bookmark_id, bookmark_name)
    cap_p.append(bm_start)

    # 2. Ajouter "Figure " (texte avant le numéro)
    r1 = OxmlElement("w:r")
    t1 = OxmlElement("w:t")
    t1.set(qn("xml:space"), "preserve")
    t1.text = f"{caption_label} "
    r1.append(t1)
    cap_p.append(r1)

    # 3. Ajouter le champ SEQ (numéro automatique) avec la bonne valeur initiale
    seq_field = _create_seq_field(seq_name, initial_value=display_num)
    cap_p.append(seq_field)

    # 4. Ajouter le bookmarkEnd
    bm_end = _create_bookmark_end(bookmark_id)
    cap_p.append(bm_end)

    # 5. Ajouter " : Légende"
    r2 = OxmlElement("w:r")
    t2 = OxmlElement("w:t")
    t2.set(qn("xml:space"), "preserve")
    t2.text = f"{nbsp}:{nbsp}{caption_text}"
    r2.append(t2)
    cap_p.append(r2)

    return bookmark_name, bookmark_id, cap_p, display_num


def _replace_text_ref_with_field(
    p_elm,
    fig_num: int,
    bookmark_name: str,
    caption_label: str = "Figure",
    display_num: int = None,
) -> int:
    """
    Remplace les occurrences de "cf. Figure N" ou "voir Figure N" dans un paragraphe
    par un champ REF pointant vers le signet.

    fig_num: Numéro original de la figure (utilisé pour la détection)
    display_num: Numéro réel à afficher (ordre d'insertion)

    Retourne le nombre de remplacements effectués.
    """
    # Utiliser display_num si fourni, sinon fig_num
    actual_display_num = display_num if display_num is not None else fig_num

    # Pattern pour détecter les références
    REF_PATTERN = re.compile(
        rf"(?i)((?:cf\.\s*|voir\s+))({re.escape(caption_label)}|fig(?:\.|ure)?|image)\s*[:\-–]?\s*{fig_num}\b"
    )

    replaced_count = 0

    try:
        t_elements = list(p_elm.xpath(".//*[local-name()='t']"))
    except Exception:
        return 0

    for t_el in t_elements:
        text = t_el.text or ""
        if not text:
            continue

        match = REF_PATTERN.search(text)
        if not match:
            continue

        # On a trouvé une référence à remplacer
        prefix = match.group(1)  # "cf. " ou "voir "

        # Récupérer le w:r parent
        r_parent = t_el.getparent()
        if r_parent is None or r_parent.tag != qn("w:r"):
            continue

        # Position dans le paragraphe
        p_parent = r_parent.getparent()
        if p_parent is None:
            continue

        r_index = list(p_parent).index(r_parent)

        # Texte avant et après la référence
        text_before = text[:match.start()]
        text_after = text[match.end():]

        # Supprimer l'ancien run
        p_parent.remove(r_parent)

        insert_pos = r_index

        # 1. Ajouter le texte avant + préfixe (cf. ou voir)
        if text_before or prefix:
            r_before = OxmlElement("w:r")
            t_before = OxmlElement("w:t")
            t_before.set(qn("xml:space"), "preserve")
            t_before.text = text_before + prefix
            r_before.append(t_before)
            p_parent.insert(insert_pos, r_before)
            insert_pos += 1

        # 2. Ajouter le champ REF avec le bon numéro d'affichage
        ref_field = _create_ref_field(bookmark_name, caption_label, actual_display_num)
        p_parent.insert(insert_pos, ref_field)
        insert_pos += 1

        # 3. Ajouter le texte après
        if text_after:
            r_after = OxmlElement("w:r")
            t_after = OxmlElement("w:t")
            t_after.set(qn("xml:space"), "preserve")
            t_after.text = text_after
            r_after.append(t_after)
            p_parent.insert(insert_pos, r_after)

        replaced_count += 1
        # On ne traite qu'une occurrence par élément t pour éviter les problèmes d'index
        break

    return replaced_count


def _replace_all_refs_in_doc(
    doc: Document,
    bookmark_map: dict,
    caption_label: str = "Figure",
) -> int:
    """
    Parcourt tout le document et remplace toutes les références textuelles
    par des champs REF.

    bookmark_map: {fig_num: (bookmark_name, display_num)}

    Retourne le nombre total de remplacements.
    """
    total_replaced = 0

    # On fait plusieurs passes car le remplacement peut fragmenter le texte
    max_passes = 5
    for _ in range(max_passes):
        pass_replaced = 0
        for p in _iter_all_paragraphs_doc(doc):
            for fig_num, (bookmark_name, display_num) in bookmark_map.items():
                count = _replace_text_ref_with_field(
                    p._p, fig_num, bookmark_name, caption_label, display_num
                )
                pass_replaced += count

        if pass_replaced == 0:
            break
        total_replaced += pass_replaced

    return total_replaced


# ==========================
# Helpers légendes
# ==========================

def _caption_for_num(captions_text, num: int):
    """
    Résout la légende pour l'image numéro ``num`` (1-indexé) à partir de ``captions_text``.
    """
    if captions_text is None:
        return None

    if isinstance(captions_text, dict):
        return captions_text.get(num)

    if isinstance(captions_text, (list, tuple)):
        return captions_text[num - 1] if 1 <= num <= len(captions_text) else None

    if isinstance(captions_text, str):
        block = captions_text.strip()

        pattern = re.compile(
            r'(?im)^\s*(?:fig(?:\.| |ure)?|image)\s*'
            r'(\d+)\s*'
            r'(?:[:\-–\.\)]+\s*)'
            r'(.+?)\s*$'
        )
        matches = pattern.findall(block)
        if matches:
            mapping: Dict[int, str] = {}
            for n_str, txt in matches:
                try:
                    idx = int(n_str)
                    if idx not in mapping:
                        mapping[idx] = txt.strip()
                except ValueError:
                    continue
            return mapping.get(num)

        lines = [ln.strip() for ln in block.splitlines() if ln.strip()]
        if len(lines) >= num and len(lines) > 1:
            return lines[num - 1]

        try:
            return captions_text.format(n=num)
        except Exception:
            return None

    return None


# ==========================
# Nettoyage des références orphelines
# ==========================

def _remove_orphan_references_from_paragraph(p_elm, orphan_nums: set[int]) -> int:
    """
    Supprime les références orphelines (cf. Figure N, voir Figure N) d'un paragraphe.
    Retourne le nombre de références supprimées.
    """
    if not orphan_nums:
        return 0

    # Pattern pour capturer les références avec contexte (parenthèses, espaces)
    # Ex: "(cf. Figure 5)", "cf. Figure 5", "(voir Figure 5)", "voir Figure 5"
    ORPHAN_RE = re.compile(
        r'\s*\(?\s*(?:cf\.\s*|voir\s+)(?:fig(?:\.|ure)?|image)\s*[:\-–]?\s*(\d+)\s*\)?\s*',
        re.IGNORECASE
    )

    removed_count = 0

    # Parcourir tous les éléments w:t (texte) du paragraphe
    try:
        t_elements = p_elm.xpath(".//*[local-name()='t']")
    except Exception:
        return 0

    for t_el in t_elements:
        text = t_el.text or ""
        if not text:
            continue

        def replace_orphan(m):
            nonlocal removed_count
            try:
                fig_num = int(m.group(1))
                if fig_num in orphan_nums:
                    removed_count += 1
                    return ""  # Supprimer la référence
            except Exception:
                pass
            return m.group(0)  # Garder la référence

        new_text = ORPHAN_RE.sub(replace_orphan, text)
        if new_text != text:
            t_el.text = new_text

    return removed_count


def _remove_orphan_references(doc: Document, orphan_nums: set[int]) -> int:
    """
    Parcourt tout le document et supprime les références orphelines.
    Retourne le nombre total de références supprimées.
    """
    if not orphan_nums:
        return 0

    total_removed = 0
    for p in _iter_all_paragraphs_doc(doc):
        total_removed += _remove_orphan_references_from_paragraph(p._p, orphan_nums)

    return total_removed


# ==========================
# Fonction principale : insertion des images par référence
# ==========================
def insert_images_by_reference_live(
    src_docx: str,
    dst_docx: str,
    out_docx: str,
    include_header_footer_src: bool = False,
    captions_text=None,
    caption_label: str = "Figure",
    caption_style: str = "Caption",
    nbspace_before_colon: bool = True,
    renumber_references: bool = False,
    target_label: str = "Figure",
    remove_orphan_references: bool = True,
    use_word_fields: bool = True,
) -> str:
    """
    Insère les images juste après le paragraphe qui contient une référence
    de type 'cf. Figure N' ou 'voir Figure N'.

    Si remove_orphan_references=True, supprime les références sans image correspondante.

    Si use_word_fields=True (défaut), utilise de vrais champs Word :
    - Légendes avec champ SEQ pour numérotation automatique
    - Signets (bookmarks) sur les légendes
    - Champs REF pour les références croisées dans le texte
    Cela permet la mise à jour automatique des numéros avec Ctrl+A puis F9.
    """
    print(
        f"[images_figures] insert_images_by_reference_live: src={src_docx}, "
        f"dst={dst_docx}, out={out_docx}"
    )

    images = extract_images_from_docx_ordered_any(
        src_docx, include_header_footer=include_header_footer_src
    )
    num_to_img = {i + 1: im for i, im in enumerate(images)}
    print(f"[images_figures] {len(num_to_img)} image(s) indexée(s) depuis le DOCX source")

    docB = Document(dst_docx)
    width_emu = _max_text_width_emu(docB)
    nbsp = "\u202F" if nbspace_before_colon else " "

    # Références uniquement (évite les légendes "Figure 1 :")
    REF_RE = re.compile(
        r"(?i)\b(?:cf\.\s*|voir\s+)(?:fig(?:\.|ure)?|image)\s*[:\-–]?\s*(\d+)\b"
    )

    inserted_for_num: set[int] = set()
    orphan_refs: set[int] = set()  # Références sans image correspondante
    inserted_count = 0
    missing_count = 0
    seen_refs: set[int] = set()
    seq_counter = 0  # Compteur d'ordre réel d'insertion
    # {fig_num: (bookmark_name, display_num)} pour les champs REF
    bookmark_map: Dict[int, tuple] = {}

    def _host_paragraph_text(p: Paragraph) -> str:
        try:
            ts = p._p.xpath(".//*[local-name()='t']")
        except Exception:
            ts = []
        txt = "".join((getattr(t, "text", None) or "") for t in ts)
        return _norm_space(txt)

    for p in _iter_all_paragraphs_doc(docB):
        txt = _host_paragraph_text(p)
        if not txt:
            continue

        nums_here = []
        for m in REF_RE.finditer(txt):
            try:
                nums_here.append(int(m.group(1)))
            except Exception:
                continue

        # dédoublonnage en gardant l’ordre
        uniq = []
        for n in nums_here:
            if n not in uniq:
                uniq.append(n)
        nums_here = uniq

        if not nums_here:
            continue

        for n in nums_here:
            seen_refs.add(n)

        # insertion après le paragraphe hôte
        for fig_num in reversed(nums_here):
            if fig_num in inserted_for_num:
                continue
            if fig_num not in num_to_img:
                missing_count += 1
                orphan_refs.add(fig_num)
                continue

            img_bytes, _ext = num_to_img[fig_num]
            new_img_p = _insert_paragraph_after_live(p._p)
            _add_picture_to_paragraph_live(docB, new_img_p, img_bytes, width_emu)

            # Incrémenter le compteur d'ordre AVANT d'ajouter la légende
            seq_counter += 1

            cap_txt = _caption_for_num(captions_text, fig_num)
            if cap_txt:
                if use_word_fields:
                    # Utiliser les vrais champs Word (SEQ + bookmark)
                    # seq_order = numéro réel d'insertion (pas le numéro IA)
                    bookmark_name, _bm_id, _cap_p, display_num = _add_caption_with_seq(
                        doc=docB,
                        after_p_elm=new_img_p,
                        fig_num=fig_num,
                        caption_text=cap_txt,
                        caption_label=caption_label,
                        caption_style=caption_style,
                        nbsp=nbsp,
                        seq_name=caption_label,
                        seq_order=seq_counter,
                    )
                    # Stocker le bookmark ET le numéro réel pour les références
                    bookmark_map[fig_num] = (bookmark_name, display_num)
                else:
                    # Ancien comportement : texte brut
                    cap_p = _insert_paragraph_after_live(new_img_p)
                    para_cap = Paragraph(cap_p, docB._body)
                    try:
                        if caption_style:
                            para_cap.style = caption_style
                    except Exception:
                        pass
                    para_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    para_cap.add_run().text = f"{caption_label} {fig_num}{nbsp}:{nbsp}{cap_txt}"

            inserted_for_num.add(fig_num)
            inserted_count += 1

    # Remplacement des références textuelles par des champs REF (si use_word_fields)
    refs_replaced_count = 0
    if use_word_fields and bookmark_map:
        refs_replaced_count = _replace_all_refs_in_doc(docB, bookmark_map, caption_label)
        print(
            f"[images_figures] Références remplacées par champs REF: {refs_replaced_count} "
            f"(figures: {sorted(bookmark_map.keys())})"
        )

    # Nettoyage des références orphelines si activé
    orphan_removed_count = 0
    if remove_orphan_references and orphan_refs:
        orphan_removed_count = _remove_orphan_references(docB, orphan_refs)
        print(
            f"[images_figures] Références orphelines supprimées: {orphan_removed_count} "
            f"(figures: {sorted(orphan_refs)})"
        )

    docB.save(out_docx)
    print(
        f"[images_figures] terminé. Références vues: {sorted(seen_refs) if seen_refs else 'aucune'}, "
        f"images insérées: {inserted_count}, références sans image: {missing_count}, "
        f"références remplacées par REF: {refs_replaced_count}, "
        f"références orphelines nettoyées: {orphan_removed_count}"
    )
    return out_docx
