# Core/writer_tpl.py
# Objectif : post-traiter le contenu généré (listes, markdown, tags ROUGE, retraits/espaces)
# SANS modifier la page de garde / titres du template, et SANS corrompre le DOCX.
#
# Principes de sécurité :
# - ne traite pas les textboxes (txbxContent) -> source fréquente de DOCX corrompu
# - ne supprime aucun paragraphe au niveau XML
# - n'insère pas de paragraphes pour les lignes vides lors du split "\n"
# - ne touche pas aux paragraphes vides (qui structurent souvent la couverture)
# - ne touche pas aux paragraphes dont le style ressemble à Title/Heading/Titre/etc.

from copy import deepcopy

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re

DEFAULT_FONT = "Times New Roman"
SPACE_RE = re.compile(r"[\s\u00A0]+")  # espaces + NBSP


def _collapse_spaces(text: str) -> str:
    if not text:
        return ""
    return SPACE_RE.sub(" ", text).strip()


def _capture_font(par):
    """
    Capture font name et size depuis le paragraphe.
    Cherche dans cet ordre :
    1) Propriétés explicites des runs
    2) Propriétés par défaut du paragraphe (pPr > rPr dans le XML)
    3) Remonte la chaîne de styles (Corps de texte → Normal → ...)
    """
    fn, fs = None, None
    # 1) Explicit run properties
    for r in par.runs:
        if r.font.name and not fn:
            fn = r.font.name
        if r.font.size is not None and fs is None:
            fs = r.font.size
        if fn and fs:
            break
    # 2) Paragraph default run properties (pPr > rPr in XML)
    if not fn or fs is None:
        try:
            pPr = par._p.find(qn('w:pPr'))
            if pPr is not None:
                rPr = pPr.find(qn('w:rPr'))
                if rPr is not None:
                    if not fn:
                        rFonts = rPr.find(qn('w:rFonts'))
                        if rFonts is not None:
                            fn = rFonts.get(qn('w:ascii')) or rFonts.get(qn('w:hAnsi'))
                    if fs is None:
                        sz = rPr.find(qn('w:sz'))
                        if sz is not None:
                            val = sz.get(qn('w:val'))
                            if val:
                                fs = Pt(int(val) // 2)
        except Exception:
            pass
    # 3) Walk up the style chain
    if not fn or fs is None:
        try:
            style = par.style
            while style is not None:
                if not fn and style.font.name:
                    fn = style.font.name
                if fs is None and style.font.size is not None:
                    fs = style.font.size
                if fn and fs:
                    break
                style = style.base_style
        except Exception:
            pass
    return fn, fs


def _apply_font(par, font_name, font_size):
    """Applique font name/size aux runs qui n'ont pas de valeur explicite."""
    for r in par.runs:
        if font_name and not r.font.name:
            r.font.name = font_name
        if font_size is not None and r.font.size is None:
            r.font.size = font_size


def _apply_markdown_style(par):
    """
    Transforme **texte** en gras et *texte* en italique dans un paragraphe.
    Gère les formats imbriqués : **gras *italique* gras** ou *italique **gras** italique*
    Reconstruit les runs en supprimant les astérisques.
    IMPORTANT : ne pas appeler sur les titres/couverture (on les skip).
    Sécurité : si les marqueurs ** ou * ne sont pas appariés, on les retire
    pour éviter de mettre tout le paragraphe en gras/italique par erreur.
    """
    raw = par.text
    if not raw or "*" not in raw:
        return

    # --- Pré-validation : retirer les marqueurs non appariés ---
    import re
    # Compter les marqueurs ** (gras) - on cherche les ** qui ne font pas partie de ***
    # D'abord remplacer temporairement *** pour ne pas les compter comme **
    temp = raw.replace("***", "\x00\x00\x00")
    bold_count = temp.count("**")
    # Compter les *** (gras+italique)
    triple_count = raw.count("***")
    if bold_count % 2 != 0:
        # Nombre impair de ** → marqueurs non appariés → retirer tous les **
        print(f"[MARKDOWN] Marqueurs ** non apparies ({bold_count}) - suppression du gras dans le paragraphe: {raw[:80]}...")
        raw = raw.replace("**", "")
    if triple_count % 2 != 0:
        raw = raw.replace("***", "")
    # Compter les * simples (italique) après nettoyage des ** et ***
    temp2 = raw.replace("**", "")
    single_star_count = temp2.count("*")
    if single_star_count % 2 != 0:
        print(f"[MARKDOWN] Marqueurs * non apparies ({single_star_count}) - suppression de l'italique dans le paragraphe: {raw[:80]}...")
        # Retirer seulement les * isolés (pas les **)
        cleaned = []
        i = 0
        while i < len(raw):
            if i + 1 < len(raw) and raw[i] == "*" and raw[i + 1] == "*":
                cleaned.append("**")
                i += 2
            elif raw[i] == "*":
                i += 1  # skip le * isolé non apparié
            else:
                cleaned.append(raw[i])
                i += 1
        raw = "".join(cleaned)

    if "*" not in raw:
        # Plus de marqueurs après nettoyage → rien à faire
        return

    par.clear()

    # Parser avec pile d'états pour gérer l'imbrication
    i = 0
    n = len(raw)
    current_text = ""
    is_bold = False
    is_italic = False

    while i < n:
        # Détection gras **
        if i + 1 < n and raw[i] == "*" and raw[i + 1] == "*":
            # Vérifier que ce n'est pas *** (gras+italique)
            if i + 2 < n and raw[i + 2] == "*":
                # *** = bascule gras ET italique
                if current_text:
                    r = par.add_run(current_text)
                    r.bold = is_bold
                    r.italic = is_italic
                    current_text = ""
                is_bold = not is_bold
                is_italic = not is_italic
                i += 3
            else:
                # ** = bascule gras
                if current_text:
                    r = par.add_run(current_text)
                    r.bold = is_bold
                    r.italic = is_italic
                    current_text = ""
                is_bold = not is_bold
                i += 2
        # Détection italique *
        elif raw[i] == "*":
            if current_text:
                r = par.add_run(current_text)
                r.bold = is_bold
                r.italic = is_italic
                current_text = ""
            is_italic = not is_italic
            i += 1
        else:
            current_text += raw[i]
            i += 1

    # Flush le reste
    if current_text:
        r = par.add_run(current_text)
        r.bold = is_bold
        r.italic = is_italic


def _iter_all_paragraphs(doc: DocxDocument):
    """
    Tous les paragraphes du document, y compris tableaux, headers/footers,
    textboxes (txbxContent) et Structured Document Tags (sdt/sdtContent).
    """
    from docx.text.paragraph import Paragraph

    def iter_container(container):
        for p in getattr(container, "paragraphs", []):
            yield p
        for table in getattr(container, "tables", []):
            for row in table.rows:
                for cell in row.cells:
                    yield from iter_container(cell)

    yield from iter_container(doc)
    for section in doc.sections:
        yield from iter_container(section.header)
        yield from iter_container(section.footer)

    # Textboxes : parcourir les paragraphes dans les txbxContent
    # parent = doc (Document) pour que .part et .style fonctionnent
    # IMPORTANT : CT_Body n'a pas .part, donc il faut passer doc (qui a .part)
    ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    body = doc.element.body
    for txbx in body.iter(f"{ns}txbxContent"):
        for p_elem in txbx.findall(f"{ns}p"):
            yield Paragraph(p_elem, doc)

    # SDT (Structured Document Tags) : certains templates rangent le contenu
    # dans des conteneurs w:sdt > w:sdtContent > w:p (ex: MEMOIRE_CII_XX.docx)
    # parent = doc (Document) pour que p.style fonctionne (résolution via doc.part)
    for sdt_content in body.iter(f"{ns}sdtContent"):
        for p_elem in sdt_content.findall(f"{ns}p"):
            yield Paragraph(p_elem, doc)
        # Tables à l'intérieur des SDT (ex: tableau résumé)
        for tbl_elem in sdt_content.findall(f"{ns}tbl"):
            for tr_elem in tbl_elem.findall(f"{ns}tr"):
                for tc_elem in tr_elem.findall(f"{ns}tc"):
                    for p_elem in tc_elem.findall(f"{ns}p"):
                        yield Paragraph(p_elem, doc)


def _apply_first_existing_style(paragraph, candidates) -> bool:
    """
    Applique le premier style existant parmi les candidats (EN/FR).
    """
    for name in candidates:
        try:
            paragraph.style = name
            return True
        except Exception:
            continue
    return False


def _ensure_bullet_numbering(doc):
    """
    S'assure qu'une définition de numérotation pour puces existe dans le document.
    Retourne le numId à utiliser.
    """
    # Vérifier si numbering.xml existe
    try:
        numbering_part = doc.part.numbering_part
        if numbering_part is None:
            # Créer numbering part si absent
            from docx.opc.constants import RELATIONSHIP_TYPE as RT
            from docx.opc.part import XmlPart
            numbering_part = doc.part.relate_to(
                XmlPart.load(doc.part.package, RT.NUMBERING, None),
                RT.NUMBERING
            )
    except Exception:
        pass

    # Retourner 1 par défaut (devrait exister dans la plupart des templates Word)
    return 1


def _apply_list_numbering(paragraph, ilvl=0, numId=1):
    """
    Applique manuellement la numérotation de liste à un paragraphe.
    ilvl: niveau de liste (0=niveau 1, 1=niveau 2)
    numId: identifiant de la numérotation (1 pour puces)
    """
    pPr = paragraph._p.get_or_add_pPr()

    # Supprimer numPr existant si présent
    for child in pPr:
        if child.tag.endswith('numPr'):
            pPr.remove(child)

    numPr = OxmlElement('w:numPr')

    # Niveau de liste
    ilvl_element = OxmlElement('w:ilvl')
    ilvl_element.set(qn('w:val'), str(ilvl))
    numPr.append(ilvl_element)

    # ID de numérotation
    numId_element = OxmlElement('w:numId')
    numId_element.set(qn('w:val'), str(numId))
    numPr.append(numId_element)

    pPr.append(numPr)


def _normalize_inline_lists(text: str) -> str:
    """
    Transforme les items de niveau 2 collés sur une ligne :
      '; / ' -> ';\n/ '
      '. / ' -> '.\n/ '
    """
    if not text:
        return ""
    s = text.replace("\r\n", "\n").replace("\r", "\n")
    s = re.sub(r"\s*;\s*/\s+", ";\n/ ", s)
    s = re.sub(r"\s*\.\s*/\s+", ".\n/ ", s)
    return s


def _style_name_lower(p) -> str:
    try:
        return (p.style.name or "").strip().lower()
    except Exception:
        return ""


# Styles qu'on ne doit PAS toucher (cover/titres/entêtes)
# Ajustez si vos templates utilisent d'autres noms.
_SKIP_STYLE_KEYWORDS = (
    "title", "titre",
    "heading", "en-tête", "entête",
    "subtitle", "sous-titre", "sous titre", "soustitre",
    "header", "footer",
    "toc", "table of contents", "table des matières",
    "caption",  # souvent utilisé pour figures
)

# Styles de CONTENU qui matchent les keywords ci-dessus mais qui doivent être traités.
# Ex: "sous-titres" (CIR) contient "titre" et "sous-titre" mais c'est un style de contenu.
_CONTENT_STYLES_WHITELIST = (
    "sous-titres",       # CIR: style de contenu pour les sections
    "corps de texte dt", # Contenu template
    "corps texte dt",    # Variante
)


def _should_skip_paragraph(p) -> bool:
    """
    Retourne True si on doit préserver intégralement ce paragraphe (page de garde, titres, etc.)
    """
    # 1) Paragraphes vides : souvent utilisés pour la mise en page de la couverture
    if (p.text or "").strip() == "":
        return True

    # 2) Styles de titres/cover (sauf whitelist de contenu)
    st = _style_name_lower(p)
    if st in _CONTENT_STYLES_WHITELIST:
        return False
    if any(k in st for k in _SKIP_STYLE_KEYWORDS):
        return True

    return False


def format_docx(path: str):
    """
    Post-traitement global :
    - tags rouge
    - éclatement multi-lignes en paragraphes (sans créer de paragraphes vides)
    - normalisation listes (-, /, 1.) en VRAIES listes Word
    - markdown léger
    - uniformisation police/espacement SUR LE CONTENU UNIQUEMENT
    - correction retraits parasites (paragraphes non-listes)
    """
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph

    doc = DocxDocument(path)

    # --- Pré-traitement CII : corriger le style des sous-titres "Description de..." ---
    # Le template CII utilise le style "DT" (ou "Sansinterligne") pour ces paragraphes,
    # mais l'utilisateur veut le style "sous-titre". On corrige ici AVANT tout traitement,
    # pour que _should_skip_paragraph les détecte et les préserve intégralement.
    _CII_SUBTITLE_PREFIXES = (
        "description du contexte",
        "description du march",
        "description détaillée de la démarche",
        "description des résultats",
    )
    for p in _iter_all_paragraphs(doc):
        st = _style_name_lower(p)
        txt_lower = (p.text or "").strip().lower()
        if txt_lower and any(txt_lower.startswith(pfx) for pfx in _CII_SUBTITLE_PREFIXES):
            if st in ("dt", "sansinterligne", "sans interligne", "no spacing"):
                try:
                    p.style = "sous-titre"
                    # Supprimer le rPr hérité du pPr (paragraph default run properties)
                    # pour laisser le style "sous-titre" contrôler le formatage (gras, etc.)
                    pPr = p._p.find(qn('w:pPr'))
                    if pPr is not None:
                        old_rPr = pPr.find(qn('w:rPr'))
                        if old_rPr is not None:
                            pPr.remove(old_rPr)
                    # Forcer le gras explicitement sur chaque run
                    # (le style lié sous-titreCar peut être ignoré par Word
                    # quand le paragraphe a déjà le style sous-titre)
                    for r in p.runs:
                        r.bold = True
                except Exception:
                    # Style "sous-titre" non trouvé → on laisse tel quel
                    pass

    # --- Supprimer le retrait négatif (hanging indent) des styles de contenu ---
    # Le style DT a w:ind left="284" hanging="284" qui crée une marge
    # à partir de la 2e ligne des paragraphes de contenu.
    # "sous-titres" (CIR) hérite de DT donc a le même problème.
    for _style_name in ("DT", "sous-titres"):
        try:
            _s = doc.styles[_style_name]
            _s.paragraph_format.left_indent = Pt(0)
            _s.paragraph_format.first_line_indent = Pt(0)
        except Exception:
            pass

    # --- Corriger la taille de police du style "Corps texte DT" ---
    # Dans le template CIR, "CorpstexteDT" définit sz=24 (12pt) au lieu de 11pt.
    # Cela affecte "Corps de texte DT" qui en hérite (présentation, partenariat, biblio).
    try:
        _s = doc.styles["Corps texte DT"]
        _s.font.size = Pt(11)
    except Exception:
        pass

    pat_a_completer = re.compile(r"(?i)(à\s*compléter\s*par\s*le\s*client\s*:?.*?)")
    pat_rien_declarer = re.compile(r"(?i)\brien\s+à\s+déclarer\b")

    # Détection listes
    order_re = re.compile(r"^(?P<num>\d+)[\.)]\s+(?P<txt>.+)$")
    bullet_lvl1_re = re.compile(r"^[\-\u2010\u2011\u2012\u2013\u2014\*]\s+(?P<txt>.+)$")
    bullet_lvl2_re = re.compile(r"^[/／]\s+(?P<txt>.+)$")
    # Chiffres romains (I, II, III, IV, V, VI, etc.) → niveau 1
    roman_re = re.compile(r"^(?P<num>[IVXLCDM]+)[\.)]\s+(?P<txt>.+)$", re.IGNORECASE)
    # Lettres (a, b, c, etc.) → niveau 2
    letter_re = re.compile(r"^(?P<letter>[a-z])[\.)]\s+(?P<txt>.+)$", re.IGNORECASE)

    def _insert_paragraph_after(par: Paragraph) -> Paragraph:
        """
        Insère un paragraphe après `par` en COPIANT les propriétés (pPr) du template,
        pour ne pas casser la mise en page (style, police, alignement).
        Supprime le gras hérité du pPr pour éviter que les nouveaux paragraphes
        héritent d'un gras non voulu du template.
        """
        new_p = OxmlElement("w:p")
        par._p.addnext(new_p)
        # Copier les propriétés de paragraphe (pPr) directement depuis le XML
        # C'est plus fiable que par.style car ça préserve tout (style, rPr par défaut, etc.)
        src_pPr = par._p.find(qn('w:pPr'))
        if src_pPr is not None:
            new_pPr = deepcopy(src_pPr)
            # Supprimer le gras hérité du template dans le pPr copié
            rPr_in_pPr = new_pPr.find(qn('w:rPr'))
            if rPr_in_pPr is not None:
                for bold_elem in rPr_in_pPr.findall(qn('w:b')):
                    rPr_in_pPr.remove(bold_elem)
                for bold_elem in rPr_in_pPr.findall(qn('w:bCs')):
                    rPr_in_pPr.remove(bold_elem)
            new_p.insert(0, new_pPr)
        new_par = Paragraph(new_p, par._parent)
        return new_par

    # 1) Éclatement des paragraphes multi-lignes (CONTENU seulement)
    paras = list(_iter_all_paragraphs(doc))
    i = 0
    while i < len(paras):
        p = paras[i]

        if _should_skip_paragraph(p):
            i += 1
            continue

        # Capturer la police ET les propriétés de paragraphe AVANT toute modification
        fn, fs = _capture_font(p)
        pPr_orig = p._p.find(qn('w:pPr'))
        pPr_split_backup = deepcopy(pPr_orig) if pPr_orig is not None else None

        raw = p.text or ""
        raw2 = _normalize_inline_lists(raw)
        if raw2 != raw:
            p.text = raw2
            # Restaurer pPr si clear() l'a supprimé
            if pPr_split_backup is not None and p._p.find(qn('w:pPr')) is None:
                p._p.insert(0, deepcopy(pPr_split_backup))
            _apply_font(p, fn, fs)

        if "\n" in (p.text or ""):
            lines = (p.text or "").replace("\r\n", "\n").replace("\r", "\n").split("\n")

            # 1ère ligne : reste dans p
            p.text = (lines[0] or "").strip()
            # Restaurer pPr si clear() l'a supprimé (AVANT _insert_paragraph_after)
            if pPr_split_backup is not None and p._p.find(qn('w:pPr')) is None:
                p._p.insert(0, deepcopy(pPr_split_backup))
            _apply_font(p, fn, fs)

            insert_after = p
            new_paras = []

            # lignes suivantes : nouveaux paragraphes (on ignore les vides)
            for ln in lines[1:]:
                ln2 = (ln or "").strip()
                if not ln2:
                    continue
                np = _insert_paragraph_after(insert_after)
                np.text = ln2
                _apply_font(np, fn, fs)
                new_paras.append(np)
                insert_after = np

            paras[i + 1 : i + 1] = new_paras

        i += 1

    # 2) Formatage contenu (tags rouge / listes / markdown / retraits / espaces)
    in_lvl2_block = False
    last_lvl2_par = None

    def close_lvl2_block():
        nonlocal in_lvl2_block, last_lvl2_par
        if in_lvl2_block and last_lvl2_par is not None:
            # dernier item lvl2 => ponctuation finale "."
            # Chercher dans tous les runs (pas seulement le dernier car il y a la puce)
            for r in reversed(last_lvl2_par.runs):
                t = r.text or ""
                # Ignorer les runs qui sont juste des puces
                if t.strip() in ["•", "○", "•\t", "○\t"]:
                    continue
                idx = t.rfind(";")
                if idx != -1:
                    r.text = t[:idx] + "." + t[idx + 1 :]
                    break
        in_lvl2_block = False
        last_lvl2_par = None

    for p in paras:
        if _should_skip_paragraph(p):
            continue

        # === Sauvegarder le formatage original du template AVANT toute modification ===
        orig_fn, orig_fs = _capture_font(p)
        pPr_elem = p._p.find(qn('w:pPr'))
        pPr_backup = deepcopy(pPr_elem) if pPr_elem is not None else None

        txt = p.text or ""

        # Tags ROUGE (sur contenu seulement)
        if pat_a_completer.search(txt):
            txt = pat_a_completer.sub(lambda m: f"[[ROUGE: {m.group(1)} ]]", txt)
        if pat_rien_declarer.search(txt):
            txt = pat_rien_declarer.sub(lambda m: f"[[ROUGE: {m.group(0)} ]]", txt)
        if txt != (p.text or ""):
            p.text = txt

        # Nettoyage indentation contenu (cause du "décalage")
        # On ne touche pas aux titres/cover car ils sont déjà skip.
        p.text = (p.text or "").replace("\t", " ").lstrip("\u00A0\u202F \t")

        stripped = (p.text or "").lstrip("\u00A0\u202F \t")
        stripped = stripped.replace("‐", "-").replace("–", "-").replace("—", "-")

        m_lvl2 = bullet_lvl2_re.match(stripped)
        m_lvl1 = bullet_lvl1_re.match(stripped)
        m_ord = order_re.match(stripped)
        m_roman = roman_re.match(stripped)
        m_letter = letter_re.match(stripped)

        is_list_item = False
        pf = p.paragraph_format

        # Lettres (a, b, c) → puces niveau 2
        if m_letter:
            content = _collapse_spaces(m_letter.group("txt")).rstrip(" ;.")
            p.clear()
            bullet_run = p.add_run("○\t")
            content_run = p.add_run(f"{content};")
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Indentation pour niveau 2
            pf.left_indent = Pt(36)  # 0.5 pouces
            pf.first_line_indent = Pt(-18)  # hanging indent
            in_lvl2_block = True
            last_lvl2_par = p
            is_list_item = True

        # Chiffres romains (I, II, III) → puces niveau 1
        elif m_roman:
            close_lvl2_block()
            content = _collapse_spaces(m_roman.group("txt"))
            p.clear()
            bullet_run = p.add_run("•\t")
            content_run = p.add_run(content)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Indentation pour niveau 1
            pf.left_indent = Pt(18)  # 0.25 pouces
            pf.first_line_indent = Pt(-18)  # hanging indent
            is_list_item = True

        elif m_lvl2:
            content = _collapse_spaces(m_lvl2.group("txt")).rstrip(" ;.")
            p.clear()
            bullet_run = p.add_run("○\t")
            content_run = p.add_run(f"{content};")
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Indentation pour niveau 2
            pf.left_indent = Pt(36)  # 0.5 pouces
            pf.first_line_indent = Pt(-18)  # hanging indent
            in_lvl2_block = True
            last_lvl2_par = p
            is_list_item = True

        elif m_lvl1:
            close_lvl2_block()
            content = _collapse_spaces(m_lvl1.group("txt"))
            p.clear()
            bullet_run = p.add_run("•\t")
            content_run = p.add_run(content)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Indentation pour niveau 1
            pf.left_indent = Pt(18)  # 0.25 pouces
            pf.first_line_indent = Pt(-18)  # hanging indent
            is_list_item = True

        elif m_ord:
            close_lvl2_block()
            content = _collapse_spaces(m_ord.group("txt"))
            p.clear()
            run = p.add_run(content)
            _apply_first_existing_style(
                p,
                ["List Number", "Liste numérotée", "List Paragraph", "Paragraphe de liste"],
            )
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            is_list_item = True

        else:
            close_lvl2_block()

        # Markdown léger (contenu seulement)
        _apply_markdown_style(p)

        # === Restaurer le formatage original du template ===
        # Restaurer pPr (style, alignement) si clear() l'a supprimé
        if pPr_backup is not None and p._p.find(qn('w:pPr')) is None:
            p._p.insert(0, pPr_backup)

        # === Supprimer le gras hérité du template dans le pPr ===
        # Le gras dans pPr > rPr fait que TOUS les runs héritent du gras par défaut,
        # même quand _apply_markdown_style() ne met pas de gras explicite.
        # Le gras doit venir UNIQUEMENT des marqueurs ** dans le texte (via _apply_markdown_style).
        pPr_current = p._p.find(qn('w:pPr'))
        if pPr_current is not None:
            rPr_in_pPr = pPr_current.find(qn('w:rPr'))
            if rPr_in_pPr is not None:
                for bold_elem in rPr_in_pPr.findall(qn('w:b')):
                    rPr_in_pPr.remove(bold_elem)
                for bold_elem in rPr_in_pPr.findall(qn('w:bCs')):
                    rPr_in_pPr.remove(bold_elem)

        # Restaurer la police d'origine sur les runs recréés
        _apply_font(p, orig_fn, orig_fs)

        # Espacements listes seulement (ne pas toucher aux paragraphes normaux)
        if is_list_item:
            pf = p.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)

    close_lvl2_block()
    doc.save(path)


def clean_custom_tags(path: str):
    """
    Remplace, dans tout le document, les séquences [[ROUGE: ...]]
    par du texte rouge, en conservant le gras / italique / underline.
    IMPORTANT : ne supprime pas de paragraphes ; uniquement des runs du paragraphe.
    Gère aussi le cas où [[ROUGE: ...]] est éclaté sur plusieurs runs XML.
    """
    doc = DocxDocument(path)

    for p in _iter_all_paragraphs(doc):
        # Ne pas toucher page de garde / titres
        if _should_skip_paragraph(p):
            continue

        full_text = "".join(r.text or "" for r in p.runs)
        if "[[ROUGE:" not in full_text:
            continue

        old_runs = list(p.runs)
        # Récupérer la mise en forme du premier run comme référence
        ref_run = old_runs[0] if old_runs else None

        # supprimer les runs existants
        for r in old_runs:
            try:
                p._p.remove(r._r)
            except Exception:
                pass

        # Traiter le texte complet (au lieu de run par run) pour gérer
        # les tags [[ROUGE:]] éclatés sur plusieurs runs XML
        parts = re.split(r"(\[\[ROUGE:.*?\]\])", full_text)

        for part in parts:
            if not part:
                continue

            def _new_run(txt, make_red=False, _ref=ref_run):
                nr = p.add_run(txt)
                if _ref is not None:
                    if _ref.font.name:
                        nr.font.name = _ref.font.name
                    nr.bold = bool(_ref.bold)
                    nr.italic = bool(_ref.italic)
                    nr.underline = bool(_ref.underline)
                    if _ref.font.size is not None:
                        nr.font.size = _ref.font.size
                if make_red:
                    nr.font.color.rgb = RGBColor(255, 0, 0)
                else:
                    if _ref is not None and _ref.font.color is not None and _ref.font.color.rgb is not None:
                        nr.font.color.rgb = _ref.font.color.rgb
                return nr

            if part.startswith("[[ROUGE:") and part.endswith("]]"):
                inner = part[len("[[ROUGE:") : -2].strip()
                if inner:
                    _new_run(inner, make_red=True)
            else:
                _new_run(part, make_red=False)

    doc.save(path)


# =============================================================================
# TABLEAU COMPARATIF CII
# =============================================================================

def insert_comparison_table(doc, placeholder: str, data: dict, client_name: str):
    """
    Insère un tableau comparatif basique (noir et blanc) à l'emplacement du placeholder.

    Args:
        doc: Document python-docx
        placeholder: Texte du placeholder (ex: "[[TABLEAU_COMPARATIF]]")
        data: Dictionnaire avec la structure:
            {
                "elements": [
                    {
                        "nom": "Element 1",
                        "client": "Oui",
                        "concurrents": {"Concurrent A": "Non", "Concurrent B": "Partiel"}
                    },
                    ...
                ]
            }
        client_name: Nom de la société cliente (pour l'en-tête)
    """
    elements = data.get("elements", [])
    if not elements:
        print("[TABLEAU] Aucun élément à insérer")
        return False

    # Collecter tous les noms de concurrents
    all_competitors = set()
    for elem in elements:
        all_competitors.update(elem.get("concurrents", {}).keys())
    competitors = sorted(all_competitors)

    if not competitors:
        print("[TABLEAU] Aucun concurrent trouvé dans les données")
        return False

    # Trouver le paragraphe contenant le placeholder (y compris dans les SDT)
    target_paragraph = None
    for para in _iter_all_paragraphs(doc):
        if placeholder in (para.text or ""):
            target_paragraph = para
            break

    if not target_paragraph:
        print(f"[TABLEAU] Placeholder '{placeholder}' non trouvé")
        return False

    # Créer le tableau
    # Colonnes: Élément | Client | Concurrent1 | Concurrent2 | ...
    num_cols = 2 + len(competitors)
    num_rows = 1 + len(elements)  # Header + lignes de données

    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = "Table Grid"

    # === En-tête (noir et blanc, texte gras centré) ===
    header_cells = table.rows[0].cells
    header_cells[0].text = "Élément"
    header_cells[1].text = client_name

    for idx, comp_name in enumerate(competitors):
        header_cells[2 + idx].text = comp_name

    # Style en-tête : gras, centré, noir
    for cell in header_cells:
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True
                run.font.name = DEFAULT_FONT
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0, 0, 0)

    # === Lignes de données ===
    for row_idx, elem in enumerate(elements):
        row = table.rows[1 + row_idx]

        # Colonne: Nom de l'élément (gras et centré)
        row.cells[0].text = elem.get("nom", "")
        for para in row.cells[0].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True
                run.font.name = DEFAULT_FONT
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0, 0, 0)

        # Colonne: Valeur client (centré)
        client_val = elem.get("client", "Oui")
        row.cells[1].text = client_val
        for para in row.cells[1].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.name = DEFAULT_FONT
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0, 0, 0)

        # Colonnes: Valeurs concurrents (centré)
        concurrents_data = elem.get("concurrents", {})
        for idx, comp_name in enumerate(competitors):
            val = concurrents_data.get(comp_name, "?")
            row.cells[2 + idx].text = val
            for para in row.cells[2 + idx].paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.name = DEFAULT_FONT
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0, 0, 0)

    # Déplacer le tableau à la position du placeholder
    tbl_element = table._tbl
    target_paragraph._element.addprevious(tbl_element)

    # Vider le paragraphe du placeholder
    target_paragraph.clear()

    print(f"[TABLEAU] Inséré avec {len(elements)} éléments et {len(competitors)} concurrents")
    return True