# app/services/templates_service.py
"""
Service pour la gestion des templates Word (CIR et CII).
Gère le stockage local et Azure Blob Storage avec historique des versions.
"""
import os
import shutil
from datetime import datetime
from typing import Tuple, List, Dict, Any

try:
    from azure.storage.blob import BlobServiceClient
except ImportError:
    BlobServiceClient = None

try:
    from docx import Document
except ImportError:
    Document = None

from app.config import settings

# Chemin local des templates
TEMPLATES_DIR = os.path.join(os.getcwd(), "Doc")
TEMPLATES_HISTORY_DIR = os.path.join(TEMPLATES_DIR, "history")

# Mapping des types de templates vers leurs fichiers
TEMPLATE_FILES = {
    "cir": "MEMOIRE_CIR2.docx",
    "cii": "MEMOIRE_CII.docx",
}

# Configuration Azure Blob Storage
BLOB_CONTAINER = "templates"


def _ensure_directories():
    """Crée les répertoires nécessaires s'ils n'existent pas."""
    os.makedirs(TEMPLATES_DIR, exist_ok=True)
    os.makedirs(TEMPLATES_HISTORY_DIR, exist_ok=True)


def _get_template_path(template_type: str) -> str:
    """Retourne le chemin complet du fichier template."""
    filename = TEMPLATE_FILES.get(template_type.lower())
    if not filename:
        raise ValueError(f"Type de template invalide: {template_type}")
    return os.path.join(TEMPLATES_DIR, filename)


def _get_history_path(template_type: str, version: int) -> str:
    """Retourne le chemin d'un fichier d'historique."""
    filename = TEMPLATE_FILES.get(template_type.lower())
    if not filename:
        raise ValueError(f"Type de template invalide: {template_type}")
    base, ext = os.path.splitext(filename)
    history_filename = f"{base}_v{version}{ext}"
    return os.path.join(TEMPLATES_HISTORY_DIR, history_filename)


def _get_blob_client():
    """Retourne un client Azure Blob Storage."""
    if not BlobServiceClient or not settings.AZURE_STORAGE_CONNECTION_STRING:
        return None
    try:
        return BlobServiceClient.from_connection_string(
            settings.AZURE_STORAGE_CONNECTION_STRING
        )
    except Exception as e:
        print(f"Erreur lors de la connexion à Azure Blob Storage: {e}")
        return None


def _upload_to_blob(template_type: str, content: bytes, version: int) -> str:
    """Upload un fichier vers Azure Blob Storage et retourne l'URL."""
    blob_client = _get_blob_client()
    if not blob_client:
        return None

    try:
        container_client = blob_client.get_container_client(BLOB_CONTAINER)
        # Créer le conteneur s'il n'existe pas
        try:
            container_client.create_container()
        except Exception:
            pass  # Le conteneur existe déjà

        filename = TEMPLATE_FILES[template_type.lower()]
        base, ext = os.path.splitext(filename)
        blob_name = f"{base}_v{version}{ext}"

        blob_client_instance = container_client.get_blob_client(blob_name)
        blob_client_instance.upload_blob(content, overwrite=True)

        return blob_client_instance.url
    except Exception as e:
        print(f"Erreur lors de l'upload vers Azure: {e}")
        return None


def get_template_info(template_type: str) -> Dict[str, Any]:
    """
    Récupère les informations d'une template.

    Returns:
        Dict avec: type, filename, size_bytes, last_modified, version, blob_url
    """
    _ensure_directories()
    template_path = _get_template_path(template_type)

    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template {template_type.upper()} non trouvée")

    stat = os.stat(template_path)

    # Déterminer la version actuelle en comptant les fichiers d'historique
    history_files = [
        f for f in os.listdir(TEMPLATES_HISTORY_DIR)
        if f.startswith(os.path.splitext(TEMPLATE_FILES[template_type.lower()])[0])
    ]
    current_version = len(history_files) + 1

    return {
        "type": template_type.lower(),
        "filename": TEMPLATE_FILES[template_type.lower()],
        "size_bytes": stat.st_size,
        "last_modified": datetime.fromtimestamp(stat.st_mtime).isoformat(),
        "version": current_version,
        "blob_url": None,  # Peut être complété avec l'URL Azure si nécessaire
    }


def get_template_content(template_type: str) -> Tuple[bytes, str]:
    """
    Lit le contenu d'une template.

    Returns:
        Tuple (content_bytes, filename)
    """
    _ensure_directories()
    template_path = _get_template_path(template_type)

    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template {template_type.upper()} non trouvée")

    with open(template_path, "rb") as f:
        content = f.read()

    return content, TEMPLATE_FILES[template_type.lower()]


def update_template(
    template_type: str,
    content: bytes,
    modified_by: str
) -> Dict[str, Any]:
    """
    Met à jour une template en créant une sauvegarde de l'ancienne version.

    Args:
        template_type: "cir" ou "cii"
        content: Contenu binaire du nouveau fichier .docx
        modified_by: Email de l'utilisateur qui fait la modification

    Returns:
        Dict avec les informations de la nouvelle version
    """
    _ensure_directories()
    template_path = _get_template_path(template_type)

    # Déterminer le numéro de version
    history_files = [
        f for f in os.listdir(TEMPLATES_HISTORY_DIR)
        if f.startswith(os.path.splitext(TEMPLATE_FILES[template_type.lower()])[0])
    ]
    new_version = len(history_files) + 1

    # Si le fichier actuel existe, le sauvegarder dans l'historique
    if os.path.exists(template_path):
        history_path = _get_history_path(template_type, new_version - 1)
        shutil.copy2(template_path, history_path)

        # Upload vers Azure Blob Storage
        with open(template_path, "rb") as f:
            old_content = f.read()
        blob_url = _upload_to_blob(template_type, old_content, new_version - 1)

    # Écrire le nouveau fichier
    with open(template_path, "wb") as f:
        f.write(content)

    # Upload de la nouvelle version vers Azure
    blob_url = _upload_to_blob(template_type, content, new_version)

    # Garder seulement les 5 dernières versions dans l'historique local
    _cleanup_old_versions(template_type)

    return {
        "version": new_version,
        "filename": TEMPLATE_FILES[template_type.lower()],
        "size_bytes": len(content),
        "modified_by": modified_by,
        "modified_at": datetime.now().isoformat(),
        "blob_url": blob_url,
    }


def _cleanup_old_versions(template_type: str):
    """Garde seulement les 5 dernières versions dans l'historique."""
    base_name = os.path.splitext(TEMPLATE_FILES[template_type.lower()])[0]
    history_files = [
        f for f in os.listdir(TEMPLATES_HISTORY_DIR)
        if f.startswith(base_name) and f.endswith(".docx")
    ]

    # Trier par numéro de version (extrait du nom de fichier)
    def extract_version(filename: str) -> int:
        try:
            # Format: template_memoire_cir_v123.docx
            version_str = filename.split("_v")[1].split(".")[0]
            return int(version_str)
        except (IndexError, ValueError):
            return 0

    history_files.sort(key=extract_version, reverse=True)

    # Supprimer les versions au-delà de la 5ème
    for old_file in history_files[5:]:
        old_path = os.path.join(TEMPLATES_HISTORY_DIR, old_file)
        try:
            os.remove(old_path)
        except Exception as e:
            print(f"Erreur lors de la suppression de {old_file}: {e}")


def list_template_history(template_type: str) -> List[Dict[str, Any]]:
    """
    Liste les 5 dernières versions d'une template.

    Returns:
        Liste de dicts avec: version, filename, size_bytes, modified_at, modified_by
    """
    _ensure_directories()
    base_name = os.path.splitext(TEMPLATE_FILES[template_type.lower()])[0]
    history_files = [
        f for f in os.listdir(TEMPLATES_HISTORY_DIR)
        if f.startswith(base_name) and f.endswith(".docx")
    ]

    def extract_version(filename: str) -> int:
        try:
            version_str = filename.split("_v")[1].split(".")[0]
            return int(version_str)
        except (IndexError, ValueError):
            return 0

    history_files.sort(key=extract_version, reverse=True)

    results = []
    for filename in history_files[:5]:
        filepath = os.path.join(TEMPLATES_HISTORY_DIR, filename)
        stat = os.stat(filepath)
        version = extract_version(filename)

        results.append({
            "version": version,
            "filename": filename,
            "size_bytes": stat.st_size,
            "modified_at": datetime.fromtimestamp(stat.st_mtime).isoformat(),
            "modified_by": None,  # Peut être enrichi avec des métadonnées
        })

    # Ajouter la version actuelle en tête
    try:
        current_info = get_template_info(template_type)
        results.insert(0, {
            "version": current_info["version"],
            "filename": current_info["filename"],
            "size_bytes": current_info["size_bytes"],
            "modified_at": current_info["last_modified"],
            "modified_by": None,
        })
    except Exception:
        pass

    return results


def restore_template_version(
    template_type: str,
    version: int,
    restored_by: str
) -> Dict[str, Any]:
    """
    Restaure une ancienne version d'une template.

    Args:
        template_type: "cir" ou "cii"
        version: Numéro de version à restaurer
        restored_by: Email de l'utilisateur qui restaure

    Returns:
        Dict avec les informations de la nouvelle version (après restauration)
    """
    _ensure_directories()
    history_path = _get_history_path(template_type, version)

    if not os.path.exists(history_path):
        raise FileNotFoundError(f"Version {version} non trouvée dans l'historique")

    # Lire le contenu de l'ancienne version
    with open(history_path, "rb") as f:
        content = f.read()

    # Utiliser update_template pour créer une nouvelle version
    return update_template(template_type, content, restored_by)


def extract_text_from_template(template_type: str) -> str:
    """
    Extrait le texte brut d'une template Word pour prévisualisation.
    Préserve les variables Jinja {variable}.

    Args:
        template_type: "cir" ou "cii"

    Returns:
        Le contenu textuel du document avec les variables Jinja
    """
    if not Document:
        raise ImportError("python-docx n'est pas installé. Installez-le avec: pip install python-docx")

    _ensure_directories()
    template_path = _get_template_path(template_type)

    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template {template_type.upper()} non trouvée")

    doc = Document(template_path)

    # Extraire tout le texte des paragraphes
    text_content = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_content.append(paragraph.text)

    # Extraire le texte des tableaux
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    text_content.append(cell_text)

    return "\n\n".join(text_content)


def update_template_variables(
    template_type: str,
    old_text: str,
    new_text: str,
    modified_by: str
) -> Dict[str, Any]:
    """
    Met à jour les variables dans une template Word en remplaçant le texte.

    Args:
        template_type: "cir" ou "cii"
        old_text: Ancien texte extrait du template
        new_text: Nouveau texte avec les variables modifiées
        modified_by: Email de l'utilisateur

    Returns:
        Dict avec les informations de la nouvelle version
    """
    if not Document:
        raise ImportError("python-docx n'est pas installé")

    _ensure_directories()
    template_path = _get_template_path(template_type)

    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template {template_type.upper()} non trouvée")

    doc = Document(template_path)

    # Créer un mapping des anciennes valeurs vers les nouvelles
    old_lines = [line.strip() for line in old_text.split("\n\n") if line.strip()]
    new_lines = [line.strip() for line in new_text.split("\n\n") if line.strip()]

    # Remplacer dans les paragraphes
    old_idx = 0
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() and old_idx < len(old_lines):
            if paragraph.text.strip() == old_lines[old_idx]:
                if old_idx < len(new_lines):
                    # Conserver le formatage mais changer le texte
                    for run in paragraph.runs:
                        run.text = ""
                    if paragraph.runs:
                        paragraph.runs[0].text = new_lines[old_idx]
                    else:
                        paragraph.text = new_lines[old_idx]
                old_idx += 1

    # Sauvegarder le document modifié
    import io
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    content = buffer.read()

    # Utiliser update_template pour créer une nouvelle version
    return update_template(template_type, content, modified_by)
